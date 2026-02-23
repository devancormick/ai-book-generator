#!/usr/bin/env python3
"""Streamlit UI for the AI Book Generator."""

import io
import json
import os
import re

import requests
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

try:
    from dotenv import load_dotenv
except ImportError:
    load_dotenv = None

if load_dotenv:
    load_dotenv()

# ── Config from .env ──────────────────────────────────────────────────────────
OPENAI_API_KEY   = os.environ.get("OPENAI_API_KEY", "")
OPENAI_MODEL     = os.environ.get("OPENAI_MODEL", "gpt-4o-mini")
GROQ_API_KEY     = os.environ.get("GROQ_API_KEY", "")
GROQ_MODEL       = os.environ.get("GROQ_MODEL", "llama-3.3-70b-versatile")
OPENROUTER_API_KEY = os.environ.get("OPENROUTER_API_KEY", "")
OPENROUTER_MODEL = os.environ.get("OPENROUTER_MODEL", "meta-llama/llama-3.1-8b-instruct:free")
OLLAMA_BASE      = os.environ.get("OLLAMA_HOST", "http://localhost:11434")
OLLAMA_MODEL     = os.environ.get("OLLAMA_MODEL", "llama3.2")


# ── Backend detection (cached for the session) ────────────────────────────────
@st.cache_resource(show_spinner="Checking available AI backends...")
def detect_backend() -> dict:
    """Try each provider in order; return info for the first valid one."""
    # OpenAI
    key = OPENAI_API_KEY.strip()
    if key and key.startswith("sk-") and len(key) >= 32:
        try:
            from openai import OpenAI
            OpenAI(api_key=key).models.list()
            return {"name": "OpenAI", "provider": "openai", "api_key": key,
                    "model": OPENAI_MODEL, "base_url": "https://api.openai.com/v1"}
        except Exception:
            pass

    # Groq
    key = GROQ_API_KEY.strip()
    if key:
        try:
            from openai import OpenAI
            OpenAI(api_key=key, base_url="https://api.groq.com/openai/v1").models.list()
            return {"name": "Groq", "provider": "groq", "api_key": key,
                    "model": GROQ_MODEL, "base_url": "https://api.groq.com/openai/v1"}
        except Exception:
            pass

    # OpenRouter
    key = OPENROUTER_API_KEY.strip()
    if key:
        try:
            from openai import OpenAI
            OpenAI(api_key=key, base_url="https://openrouter.ai/api/v1").models.list()
            return {"name": "OpenRouter", "provider": "openrouter", "api_key": key,
                    "model": OPENROUTER_MODEL, "base_url": "https://openrouter.ai/api/v1"}
        except Exception:
            pass

    # Ollama
    try:
        r = requests.get(f"{OLLAMA_BASE.rstrip('/')}/api/tags", timeout=5)
        r.raise_for_status()
        return {"name": f"Ollama ({OLLAMA_MODEL})", "provider": "ollama",
                "api_key": None, "model": OLLAMA_MODEL, "base_url": OLLAMA_BASE}
    except Exception:
        pass

    return {"name": "None", "provider": "none", "api_key": None, "model": None, "base_url": None}


def call_llm(system: str, user: str, backend: dict) -> str:
    """Send a prompt to the active backend and return the response text."""
    if backend["provider"] == "ollama":
        url = f"{backend['base_url'].rstrip('/')}/api/generate"
        r = requests.post(url, json={
            "model": backend["model"],
            "prompt": f"{system}\n\n{user}",
            "stream": False,
            "options": {"num_predict": 8192},
        }, timeout=600)
        r.raise_for_status()
        return (r.json().get("response") or "").strip()

    # OpenAI-compatible (OpenAI / Groq / OpenRouter)
    from openai import OpenAI
    extra = {}
    if backend["provider"] == "openrouter":
        extra["default_headers"] = {
            "HTTP-Referer": "https://github.com/devancormick/ai-book-generator",
            "X-Title": "AI Book Generator",
        }
    client = OpenAI(api_key=backend["api_key"], base_url=backend["base_url"], **extra)
    resp = client.chat.completions.create(
        model=backend["model"],
        messages=[{"role": "system", "content": system}, {"role": "user", "content": user}],
        max_tokens=8192,
    )
    return (resp.choices[0].message.content or "").strip()


# ── Generation functions ──────────────────────────────────────────────────────
def generate_description(topic: str, backend: dict) -> str:
    return call_llm(
        "You are a creative fiction writer. Be concise.",
        f"Write a 3-4 sentence murder mystery premise for the topic: '{topic}'. "
        "Include: setting, victim, protagonist/detective, and the central mystery. "
        "Output only the description, no preamble.",
        backend,
    )


def generate_outline(topic: str, description: str, num_chapters: int, pov: str, backend: dict) -> list[dict]:
    pov_note = "first-person (detective narrates using 'I')" if pov == "first" else "third person"
    raw = call_llm(
        "You are a professional fiction writer. Output only valid JSON, no markdown.",
        f"Create a detailed murder mystery outline with exactly {num_chapters} chapters.\n"
        f"Topic: {topic}\nPremise: {description}\nPOV: {pov_note}\n"
        "Output a JSON array. First object MUST include \"novel_title\" (unique, specific title). "
        "Each object: \"chapter_number\", \"title\", \"summary\" (2-4 sentences).",
        backend,
    )
    raw = re.sub(r"^```\w*\n?", "", raw)
    raw = re.sub(r"\n?```\s*$", "", raw)
    try:
        outline = json.loads(raw)
        if isinstance(outline, list) and len(outline) >= num_chapters:
            return outline[:num_chapters]
    except json.JSONDecodeError:
        pass
    # Fallback
    items = [{"chapter_number": i, "title": f"Chapter {i}", "summary": f"Events of chapter {i}."}
             for i in range(1, num_chapters + 1)]
    items[0]["novel_title"] = topic
    return items


def _build_state_summary(state: dict) -> str:
    parts = []
    if state.get("characters"):
        parts.append("Characters: " + "; ".join(state["characters"]))
    if state.get("locations"):
        parts.append("Locations: " + "; ".join(state["locations"]))
    if state.get("clues"):
        parts.append("Clues/events: " + "; ".join(state["clues"]))
    if state.get("recent_summaries"):
        parts.append("Recent summaries:\n" + "\n".join(state["recent_summaries"][-3:]))
    return "\n\n".join(parts) or "No prior state."


def generate_chapter(spec: dict, outline: list[dict], state: dict,
                     ch_num: int, pov: str, backend: dict) -> str:
    pov_instr = ("first person past tense (narrator is the detective/protagonist, using 'I')"
                 if pov == "first" else "third person past tense")
    outline_text = "\n".join(
        f"Ch{n.get('chapter_number', i)}: {n.get('title', '')} - {n.get('summary', '')}"
        for i, n in enumerate(outline, 1)
    )
    return call_llm(
        f"You are a professional murder mystery writer. Write in {pov_instr}. "
        "Maintain continuity with the outline and story state. "
        "Aim for approximately 2500 words. Output only chapter prose, no title or chapter number.",
        f"Outline:\n{outline_text}\n\nStory state:\n{_build_state_summary(state)}\n\n"
        f"Write chapter {ch_num}: {spec.get('title', f'Chapter {ch_num}')}.\n"
        f"Summary: {spec.get('summary', '')}",
        backend,
    )


def update_state(text: str, state: dict, backend: dict) -> dict:
    state = {k: list(v) for k, v in state.items()}
    for k in ("characters", "locations", "clues", "recent_summaries"):
        state.setdefault(k, [])
    try:
        raw = call_llm(
            "Extract from the chapter: characters (name-role), locations, key clues/events. "
            "Output JSON only: {\"characters\": [], \"locations\": [], \"clues\": []}",
            f"Extract from:\n\n{text[:8000]}",
            backend,
        )
        raw = re.sub(r"^```\w*\n?", "", raw)
        raw = re.sub(r"\n?```\s*$", "", raw)
        data = json.loads(raw)
        for key in ("characters", "locations", "clues"):
            for item in data.get(key, []):
                if item and item not in state[key]:
                    state[key].append(item)
    except Exception:
        pass
    return state


def summarize_chapter(text: str, backend: dict) -> str:
    return call_llm(
        "Summarize this chapter in 2-4 sentences for continuity. No preamble.",
        f"Summarize:\n\n{text[:6000]}",
        backend,
    )


# ── Export helpers ────────────────────────────────────────────────────────────
def build_md(chapters: list[tuple[str, str]], title: str, author: str) -> str:
    lines = [f"# {title}\n\n"]
    lines.append(f"*By {author}*\n\n" if author else "*Generated murder mystery manuscript.*\n\n")
    for ch_title, body in chapters:
        lines.append(f"## {ch_title}\n\n{body.strip()}\n\n---\n\n")
    return "".join(lines)


def build_docx(chapters: list[tuple[str, str]], title: str, author: str) -> bytes:
    doc = Document()
    h = doc.add_heading(title, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    byline = doc.add_paragraph(f"By {author}" if author else "Generated murder mystery manuscript.")
    byline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    byline.paragraph_format.space_after = Pt(24)
    for i, (ch_title, body) in enumerate(chapters):
        if i > 0:
            doc.add_page_break()
        doc.add_heading(ch_title, level=1)
        for para in body.split("\n\n"):
            para = para.strip()
            if para:
                p = doc.add_paragraph(para)
                p.paragraph_format.space_after = Pt(12)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _safe_filename(title: str) -> str:
    return re.sub(r"[^\w\s-]", "", title).strip().replace(" ", "_")[:60]


# ── Streamlit page ────────────────────────────────────────────────────────────
st.set_page_config(page_title="AI Book Generator", page_icon="book", layout="wide")

st.title("AI Book Generator")
st.caption("Generate a full murder mystery novel — chapter by chapter, with AI.")

# Session state defaults
for key, default in [("description", ""), ("chapters_out", []), ("novel_title", "")]:
    if key not in st.session_state:
        st.session_state[key] = default

# ── Sidebar ───────────────────────────────────────────────────────────────────
with st.sidebar:
    st.header("Settings")
    num_chapters = st.selectbox("Number of chapters", [10, 11, 12], index=1)
    pov          = st.radio("Point of view", ["first", "third"],
                             format_func=lambda x: f"{x.capitalize()} person")
    author_name  = st.text_input("Author name (optional)")
    st.divider()
    st.subheader("Backend")
    backend = detect_backend()
    if backend["provider"] == "none":
        st.error("No backend available. Set a key in .env or start Ollama.")
    else:
        st.success(f"{backend['name']}\n{backend['model']}")

# ── Inputs ────────────────────────────────────────────────────────────────────
topic = st.text_input(
    "Topic / Title idea",
    placeholder="e.g. A murder at a Victorian manor during a snowstorm",
    max_chars=200,
)

col_area, col_btn = st.columns([5, 1])
with col_area:
    description = st.text_area(
        "Premise / Description",
        value=st.session_state.description,
        placeholder="Describe the setting, characters, and mystery — or click Generate to auto-fill.",
        height=130,
    )
with col_btn:
    st.write("")
    st.write("")
    st.write("")
    if st.button("Generate\nDescription", disabled=not topic or backend["provider"] == "none",
                 use_container_width=True, help="Auto-generate a premise from the topic"):
        with st.spinner("Writing premise..."):
            st.session_state.description = generate_description(topic, backend)
        st.rerun()

st.divider()

# ── Generate button ───────────────────────────────────────────────────────────
can_generate = bool(topic) and backend["provider"] != "none"
if st.button("Generate Novel", type="primary", disabled=not can_generate, use_container_width=False):
    desc = description.strip() or st.session_state.description or topic
    st.session_state.chapters_out = []
    st.session_state.novel_title  = ""

    progress   = st.progress(0, text="Generating outline...")
    status_box = st.empty()
    log_area   = st.container()

    # Outline
    with st.spinner("Generating outline..."):
        outline = generate_outline(topic, desc, num_chapters, pov, backend)

    novel_title = ""
    for item in outline:
        if isinstance(item, dict):
            t = str(item.get("novel_title", "")).strip()
            if t:
                novel_title = t
                break
    if not novel_title:
        novel_title = topic
    st.session_state.novel_title = novel_title
    status_box.info(f'Outline ready: **"{novel_title}"** ({len(outline)} chapters)')

    # Chapter loop
    story_state   = {"characters": [], "locations": [], "clues": [], "recent_summaries": []}
    chapters_out  = []

    for i, spec in enumerate(outline):
        ch_num  = i + 1
        ch_title = spec.get("title", f"Chapter {ch_num}")
        progress.progress(i / len(outline),
                          text=f"Writing chapter {ch_num}/{len(outline)}: {ch_title}...")

        text = generate_chapter(spec, outline, story_state, ch_num, pov, backend)

        # Extend if too short
        if len(text.split()) < 1500:
            extra = call_llm(
                f"Continue this chapter in the same style. Add ~{1500 - len(text.split())} words. Output only the continuation.",
                f"Continue:\n\n{text[-1500:]}",
                backend,
            )
            text += "\n\n" + extra

        chapters_out.append((ch_title, text))

        summary = summarize_chapter(text, backend)
        story_state["recent_summaries"].append(f"Ch{ch_num}: {summary}")
        story_state = update_state(text, story_state, backend)

        # Show each completed chapter inline
        with log_area:
            with st.expander(f"Chapter {ch_num}: {ch_title}  (~{len(text.split()):,} words)",
                             expanded=False):
                st.markdown(text[:3000] + ("\n\n*...preview truncated*" if len(text) > 3000 else ""))

    st.session_state.chapters_out = chapters_out
    total_words = sum(len(b.split()) for _, b in chapters_out)
    progress.progress(1.0, text="Done!")
    status_box.success(
        f'Novel complete!  **"{novel_title}"**  —  {len(chapters_out)} chapters  ·  ~{total_words:,} words'
    )

# ── Download section ──────────────────────────────────────────────────────────
if st.session_state.chapters_out:
    title    = st.session_state.novel_title or topic or "Murder_Mystery"
    chapters = st.session_state.chapters_out
    fname    = _safe_filename(title)

    st.divider()
    st.subheader("Download manuscript")

    col_md, col_docx, col_stats = st.columns([1, 1, 2])

    with col_md:
        md_bytes = build_md(chapters, title, author_name).encode("utf-8")
        st.download_button(
            "Download .md",
            data=md_bytes,
            file_name=f"{fname}.md",
            mime="text/markdown",
            use_container_width=True,
        )

    with col_docx:
        docx_bytes = build_docx(chapters, title, author_name)
        st.download_button(
            "Download .docx",
            data=docx_bytes,
            file_name=f"{fname}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    with col_stats:
        total_words = sum(len(b.split()) for _, b in chapters)
        st.caption(
            f"**{len(chapters)}** chapters  ·  "
            f"**~{total_words:,}** words  ·  "
            f"Backend: {backend['name']}"
        )
