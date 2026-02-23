#!/usr/bin/env python3
"""
AI Book Generator – local script to generate a murder mystery novel (20k–30k words)
using OpenAI, Groq, OpenRouter, or local Llama via Ollama. Outputs .docx (default) or .md.
"""

import argparse
import json
import os
import re
import sys
import time
from datetime import datetime
from pathlib import Path
from typing import Optional

try:
    from dotenv import load_dotenv
except ImportError:
    load_dotenv = None

if load_dotenv:
    load_dotenv()

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")
OPENAI_MODEL = os.environ.get("OPENAI_MODEL", "gpt-4o-mini")

GROQ_API_KEY = os.environ.get("GROQ_API_KEY")
GROQ_MODEL = os.environ.get("GROQ_MODEL", "llama-3.3-70b-versatile")

OPENROUTER_API_KEY = os.environ.get("OPENROUTER_API_KEY")
OPENROUTER_MODEL = os.environ.get("OPENROUTER_MODEL", "meta-llama/llama-3.1-8b-instruct:free")

OLLAMA_BASE = os.environ.get("OLLAMA_HOST", "http://localhost:11434")
OLLAMA_MODEL = os.environ.get("OLLAMA_MODEL", "llama3.2")


def _parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate a murder mystery novel (20k–30k words).")
    parser.add_argument(
        "--format", "-f",
        choices=["md", "docx"],
        default=None,
        help="Output format: md (default) or docx",
    )
    parser.add_argument(
        "--chapters",
        type=int,
        default=None,
        choices=range(10, 13),
        metavar="10-12",
        help="Number of chapters (10–12 for full novel; default 11)",
    )
    parser.add_argument(
        "--test",
        action="store_true",
        default=False,
        help="Quick test run: 2 short chapters (~800 words each)",
    )
    parser.add_argument(
        "--output",
        default=None,
        metavar="DIR",
        help="Directory to save the manuscript (default: ./output/)",
    )
    parser.add_argument(
        "--author",
        default=None,
        metavar="NAME",
        help="Author name shown on the manuscript byline",
    )
    parser.add_argument(
        "--model",
        default=None,
        metavar="MODEL",
        help="Override the LLM model name (OpenAI model ID or Ollama model name)",
    )
    parser.add_argument(
        "--pov",
        choices=["first", "third"],
        default="first",
        help="Narrative point of view: first (default) or third person",
    )
    args, _ = parser.parse_known_args()
    return args


_ARGS = _parse_args()

# TEST_RUN: --test flag takes precedence, then env var
TEST_RUN = _ARGS.test or os.environ.get("TEST_RUN", "").strip() in ("1", "true", "yes")

# --model overrides the model name for whichever backend is active
if _ARGS.model:
    OPENAI_MODEL = _ARGS.model
    GROQ_MODEL = _ARGS.model
    OPENROUTER_MODEL = _ARGS.model
    OLLAMA_MODEL = _ARGS.model

# Narrative point-of-view setting
_POV = _ARGS.pov  # "first" or "third"
POV_INSTRUCTION = (
    "first person past tense (the narrator is the detective/protagonist, using 'I')"
    if _POV == "first"
    else "third person past tense"
)


def _parse_format() -> str:
    if _ARGS.format:
        return _ARGS.format
    raw = (os.environ.get("OUTPUT_FORMAT", "md") or "md").strip().lower()
    return "docx" if raw == "docx" else "md"


OUTPUT_FORMAT = _parse_format()
_ch = _ARGS.chapters if _ARGS.chapters is not None else None
_env_ch = os.environ.get("TARGET_CHAPTERS")
TARGET_CHAPTERS = 2 if TEST_RUN else (_ch or (int(_env_ch) if _env_ch else 11))
if not TEST_RUN and TARGET_CHAPTERS not in (10, 11, 12):
    TARGET_CHAPTERS = 11
TARGET_WORDS_PER_CHAPTER = 800 if TEST_RUN else 2500
MIN_CHAPTER_WORDS_RATIO = 0.6
GENRE = "murder mystery"

_MAX_RETRIES = 3
_RETRY_DELAY = 5  # seconds between retry attempts

# ── Per-provider validity cache ───────────────────────────────────────────────
_OPENAI_VALID: Optional[bool] = None
_GROQ_VALID: Optional[bool] = None
_OPENROUTER_VALID: Optional[bool] = None


def _validate_openai_key() -> bool:
    global _OPENAI_VALID
    if _OPENAI_VALID is not None:
        return _OPENAI_VALID
    key = (OPENAI_API_KEY or "").strip()
    if not key or not key.startswith("sk-") or len(key) < 32:
        _OPENAI_VALID = False
        return False
    try:
        from openai import OpenAI
        OpenAI(api_key=key).models.list()
        _OPENAI_VALID = True
        return True
    except Exception:
        _OPENAI_VALID = False
        return False


def _validate_groq_key() -> bool:
    global _GROQ_VALID
    if _GROQ_VALID is not None:
        return _GROQ_VALID
    key = (GROQ_API_KEY or "").strip()
    if not key:
        _GROQ_VALID = False
        return False
    try:
        from openai import OpenAI
        OpenAI(api_key=key, base_url="https://api.groq.com/openai/v1").models.list()
        _GROQ_VALID = True
        return True
    except Exception:
        _GROQ_VALID = False
        return False


def _validate_openrouter_key() -> bool:
    global _OPENROUTER_VALID
    if _OPENROUTER_VALID is not None:
        return _OPENROUTER_VALID
    key = (OPENROUTER_API_KEY or "").strip()
    if not key:
        _OPENROUTER_VALID = False
        return False
    try:
        from openai import OpenAI
        OpenAI(api_key=key, base_url="https://openrouter.ai/api/v1").models.list()
        _OPENROUTER_VALID = True
        return True
    except Exception:
        _OPENROUTER_VALID = False
        return False


def use_openai() -> bool:
    return _validate_openai_key()

def use_groq() -> bool:
    return not use_openai() and _validate_groq_key()

def use_openrouter() -> bool:
    return not use_openai() and not _validate_groq_key() and _validate_openrouter_key()


def _active_backend() -> str:
    """Return a human-readable name for the backend that will be used."""
    if use_openai():
        return f"OpenAI ({OPENAI_MODEL})"
    if _validate_groq_key():
        return f"Groq ({GROQ_MODEL})"
    if _validate_openrouter_key():
        return f"OpenRouter ({OPENROUTER_MODEL})"
    return f"Ollama/{OLLAMA_MODEL} at {OLLAMA_BASE}"


def check_backend():
    """Verify the chosen backend is reachable before generation starts."""
    if use_openai() or _validate_groq_key() or _validate_openrouter_key():
        return
    # Ollama fallback — check it is running and the model is pulled
    try:
        r = requests.get(f"{OLLAMA_BASE.rstrip('/')}/api/tags", timeout=5)
        r.raise_for_status()
        tags = r.json()
        available = [m.get("name", "") for m in tags.get("models", [])]
        model_base = OLLAMA_MODEL.split(":")[0]
        if available and not any(model_base in m for m in available):
            print(
                f"Warning: model '{OLLAMA_MODEL}' not found in Ollama. "
                f"Available: {', '.join(available) or 'none'}",
                file=sys.stderr,
            )
            print(f"  Pull it first with: ollama pull {OLLAMA_MODEL}", file=sys.stderr)
    except requests.exceptions.RequestException as e:
        print("No cloud API key set and Ollama is not reachable.", file=sys.stderr)
        print(f"  {e}", file=sys.stderr)
        print(
            "  Options:\n"
            "    • Set GROQ_API_KEY in .env for free cloud inference (recommended)\n"
            "    • Set OPENROUTER_API_KEY in .env for free cloud inference\n"
            "    • Start Ollama locally: ollama serve",
            file=sys.stderr,
        )
        sys.exit(1)


def _complete_compat(base_url: str, api_key: str, model: str, system: str, user: str) -> str:
    """Call any OpenAI-compatible API (OpenAI, Groq, OpenRouter)."""
    from openai import OpenAI
    extra = {}
    if "openrouter" in base_url:
        extra["default_headers"] = {
            "HTTP-Referer": "https://github.com/devancormick/ai-book-generator",
            "X-Title": "AI Book Generator",
        }
    client = OpenAI(api_key=api_key, base_url=base_url, **extra)
    resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        max_tokens=8192,
    )
    return (resp.choices[0].message.content or "").strip()


def complete_openai(system: str, user: str) -> str:
    return _complete_compat("https://api.openai.com/v1", OPENAI_API_KEY, OPENAI_MODEL, system, user)

def complete_groq(system: str, user: str) -> str:
    return _complete_compat("https://api.groq.com/openai/v1", GROQ_API_KEY, GROQ_MODEL, system, user)

def complete_openrouter(system: str, user: str) -> str:
    return _complete_compat("https://openrouter.ai/api/v1", OPENROUTER_API_KEY, OPENROUTER_MODEL, system, user)


def complete_ollama(system: str, user: str) -> str:
    url = f"{OLLAMA_BASE.rstrip('/')}/api/generate"
    payload = {
        "model": OLLAMA_MODEL,
        "prompt": f"{system}\n\n{user}",
        "stream": False,
        "options": {"num_predict": 8192},
    }
    r = requests.post(url, json=payload, timeout=600)
    r.raise_for_status()
    data = r.json()
    return (data.get("response") or "").strip()


def _is_quota_error(e: Exception) -> bool:
    """Return True for permanent billing/quota errors — do not retry, try next provider."""
    msg = str(e).lower()
    return "insufficient_quota" in msg or "billing" in msg or "rate_limit" not in msg and "429" in str(e)


def complete(system: str, user: str) -> str:
    """Call the best available backend in priority order: OpenAI → Groq → OpenRouter → Ollama.
    On permanent quota/billing errors the exhausted provider is skipped automatically."""
    global _OPENAI_VALID, _GROQ_VALID, _OPENROUTER_VALID

    providers = [
        ("OpenAI",      use_openai,          complete_openai,      lambda: setattr(__builtins__, '_', None) or globals().update(_OPENAI_VALID=False)),
        ("Groq",        _validate_groq_key,  complete_groq,        lambda: globals().update(_GROQ_VALID=False)),
        ("OpenRouter",  _validate_openrouter_key, complete_openrouter, lambda: globals().update(_OPENROUTER_VALID=False)),
        ("Ollama",      lambda: True,        complete_ollama,      lambda: None),
    ]

    last_err: Optional[Exception] = None
    for attempt in range(1, _MAX_RETRIES + 1):
        for name, is_active, call_fn, invalidate in providers:
            if not is_active():
                continue
            try:
                return call_fn(system, user)
            except Exception as e:
                last_err = e
                if _is_quota_error(e):
                    print(f"  {name} quota/billing error — skipping to next provider.", file=sys.stderr)
                    if name == "OpenAI":
                        _OPENAI_VALID = False
                    elif name == "Groq":
                        _GROQ_VALID = False
                    elif name == "OpenRouter":
                        _OPENROUTER_VALID = False
                    break  # try next provider immediately
                # Transient error — retry this provider after delay
                if attempt < _MAX_RETRIES:
                    print(
                        f"  {name} call failed (attempt {attempt}/{_MAX_RETRIES}): {e}. "
                        f"Retrying in {_RETRY_DELAY}s...",
                        file=sys.stderr,
                    )
                    time.sleep(_RETRY_DELAY)
                break
    raise RuntimeError(f"All LLM backends failed after {_MAX_RETRIES} attempts. Last error: {last_err}") from last_err


def generate_outline() -> list[dict]:
    system = (
        "You are a professional fiction writer. Output only valid JSON, no markdown or extra text."
    )
    user = (
        f"Create a detailed outline for a {GENRE} novel with exactly {TARGET_CHAPTERS} chapters. "
        f"Each chapter should be about {TARGET_WORDS_PER_CHAPTER} words. "
        "Include: title, setting, main characters (suspects and victim), central mystery, and red herrings. "
        "Output a JSON array of objects. First object MUST include \"novel_title\": a unique, specific book title "
        "(e.g. \"The Ashford Manor Affair\", \"Death at the Winter Ball\")—never generic like \"Murder Mystery\". "
        "Each object: \"chapter_number\" (1-based), \"title\", \"summary\" (2-4 sentences)."
    )
    raw = complete(system, user)
    raw = re.sub(r"^```\w*\n?", "", raw)
    raw = re.sub(r"\n?```\s*$", "", raw)
    try:
        outline = json.loads(raw)
    except json.JSONDecodeError:
        outline = _fallback_outline()
    if not isinstance(outline, list) or len(outline) < TARGET_CHAPTERS:
        outline = _fallback_outline()
    return outline[:TARGET_CHAPTERS]


def _fallback_outline() -> list[dict]:
    items = [
        {"chapter_number": i, "title": f"Chapter {i}", "summary": f"Key events for chapter {i}."}
        for i in range(1, TARGET_CHAPTERS + 1)
    ]
    if items:
        items[0]["novel_title"] = "The Shadow at Thornwood Hall"
    return items


def build_story_state_summary(state: dict) -> str:
    parts = []
    if state.get("characters"):
        parts.append("Characters: " + "; ".join(state["characters"]))
    if state.get("locations"):
        parts.append("Locations: " + "; ".join(state["locations"]))
    if state.get("clues"):
        parts.append("Clues/events: " + "; ".join(state["clues"]))
    if state.get("recent_summaries"):
        parts.append("Recent chapter summaries:\n" + "\n".join(state["recent_summaries"][-3:]))
    return "\n\n".join(parts) if parts else "No prior story state."


def generate_chapter(
    chapter_spec: dict,
    full_outline: list[dict],
    story_state: dict,
    chapter_index: int,
) -> str:
    outline_text = "\n".join(
        f"Ch{n.get('chapter_number', i)}: {n.get('title', '')} – {n.get('summary', '')}"
        for i, n in enumerate(full_outline, 1)
    )
    state_summary = build_story_state_summary(story_state)

    system = (
        f"You are a professional writer of {GENRE} fiction. "
        f"Write in {POV_INSTRUCTION}. Maintain continuity with the given outline and story state. "
        f"Aim for approximately {TARGET_WORDS_PER_CHAPTER} words for this chapter. "
        "Output only the chapter prose, no title or chapter number."
    )
    user = (
        f"Outline:\n{outline_text}\n\n"
        f"Story state so far:\n{state_summary}\n\n"
        f"Write chapter {chapter_index}: {chapter_spec.get('title', 'Chapter ' + str(chapter_index))}. "
        f"Summary to follow: {chapter_spec.get('summary', '')}"
    )
    return complete(system, user)


def extract_state_updates(text: str, existing: dict) -> dict:
    """Extract characters, locations, clues from chapter via LLM and merge into state."""
    existing = dict(existing)
    for k in ("characters", "locations", "clues", "recent_summaries"):
        existing.setdefault(k, [])
    system = (
        "You are an editor. Output only valid JSON, no markdown. "
        "Extract from the chapter: characters (names/roles), locations, clues/plot events. "
        "Output: {\"characters\": [\"name - role\"], \"locations\": [\"place\"], \"clues\": [\"event or clue\"]}"
    )
    # Use up to 8000 chars so we cover the full chapter (a 2500-word chapter is ~16k chars)
    user = f"Extract from this murder mystery chapter:\n\n{text[:8000]}"
    try:
        raw = complete(system, user)
        raw = re.sub(r"^```\w*\n?", "", raw)
        raw = re.sub(r"\n?```\s*$", "", raw)
        data = json.loads(raw)
        for name in data.get("characters", []):
            if name and name not in existing["characters"]:
                existing["characters"].append(name)
        for loc in data.get("locations", []):
            if loc and loc not in existing["locations"]:
                existing["locations"].append(loc)
        for clue in data.get("clues", []):
            if clue and clue not in existing["clues"]:
                existing["clues"].append(clue)
    except (json.JSONDecodeError, KeyError):
        pass
    return existing


def summarize_chapter(chapter_text: str, story_state: dict) -> str:
    system = "You are an editor. Output only a short summary (2-4 sentences) of the given chapter for continuity. No preamble."
    # Use up to 6000 chars for better summarization coverage of longer chapters
    user = f"Summarize this chapter for continuity:\n\n{chapter_text[:6000]}"
    return complete(system, user)


def export_md(
    chapters: list[tuple[str, str]],
    output_path: Path,
    title: str = "Murder Mystery",
    author: str = "",
):
    lines = [f"# {title}\n\n"]
    if author:
        lines.append(f"*By {author}*\n\n")
    else:
        lines.append("*Generated manuscript – murder mystery.*\n\n")
    for chapter_title, body in chapters:
        lines.append(f"## {chapter_title}\n\n")
        lines.append(body.strip())
        lines.append("\n\n---\n\n")
    output_path.write_text("".join(lines), encoding="utf-8")


def export_docx(
    chapters: list[tuple[str, str]],
    output_path: Path,
    title: str = "Murder Mystery",
    author: str = "",
):
    doc = Document()

    # Centered title
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Centered byline
    byline_text = f"By {author}" if author else "Generated manuscript – murder mystery."
    byline = doc.add_paragraph(byline_text)
    byline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    byline.paragraph_format.space_after = Pt(24)

    for i, (chapter_title, body) in enumerate(chapters):
        # Page break before every chapter except the first
        if i > 0:
            doc.add_page_break()
        doc.add_heading(chapter_title, level=1)
        for para in body.split("\n\n"):
            para = para.strip()
            if not para:
                continue
            p = doc.add_paragraph(para)
            p.paragraph_format.space_after = Pt(12)

    doc.save(output_path)


def main():
    print("AI Book Generator – murder mystery (local script)")
    if TEST_RUN:
        print("Test run: 2 short chapters (~800 words each).")
    print(f"Output format: {OUTPUT_FORMAT}")
    print(f"Point of view: {_POV} person")
    print(f"Backend: {_active_backend()}")
    check_backend()

    print("Generating outline...")
    outline = generate_outline()
    novel_title = ""
    title_slug = ""
    for item in (outline or []):
        if isinstance(item, dict):
            raw_title = str(item.get("novel_title", "")).strip()
            if raw_title:
                novel_title = raw_title
                title_slug = "_" + re.sub(r"[^\w\s-]", "", raw_title).strip().replace(" ", "_")[:50]
            elif not novel_title and item.get("title"):
                novel_title = str(item["title"]).strip()
                title_slug = "_" + re.sub(r"[^\w\s-]", "", novel_title).strip().replace(" ", "_")[:50] if novel_title else ""
            break
    if not novel_title:
        novel_title = outline[0].get("title", "Untitled") if outline and isinstance(outline[0], dict) else "Untitled"
        title_slug = "_" + re.sub(r"[^\w\s-]", "", novel_title).strip().replace(" ", "_")[:50] if novel_title else ""
    print(f'Outline: {len(outline)} chapters — "{novel_title}"')

    story_state: dict = {
        "characters": [],
        "locations": [],
        "clues": [],
        "recent_summaries": [],
    }
    chapters_out: list[tuple[str, str]] = []
    author = _ARGS.author or ""

    for i, spec in enumerate(outline):
        ch_num = i + 1
        ch_title = spec.get("title", f"Chapter {ch_num}")
        print(f"Writing chapter {ch_num}/{len(outline)}: {ch_title}...")
        text = generate_chapter(spec, outline, story_state, ch_num)
        min_words = int(TARGET_WORDS_PER_CHAPTER * MIN_CHAPTER_WORDS_RATIO)
        if len(text.split()) < min_words:
            print(f"  Chapter too short ({len(text.split())} words), extending...")
            ext_system = (
                f"You are a writer. Continue this chapter in the same style, "
                f"writing in {POV_INSTRUCTION}. "
                f"Add approximately {min_words - len(text.split())} more words. "
                "Output only the continuation, no recap."
            )
            ext_user = f"Continue from here:\n\n{text[-1500:]}"
            text += "\n\n" + complete(ext_system, ext_user)
        ch_words = len(text.split())
        print(f"  Done: ~{ch_words:,} words.")
        chapters_out.append((ch_title, text))
        summary = summarize_chapter(text, story_state)
        story_state["recent_summaries"].append(f"Ch{ch_num}: {summary}")
        story_state = extract_state_updates(text, story_state)

    ext = "docx" if OUTPUT_FORMAT == "docx" else "md"
    base = "manuscript_test" if TEST_RUN else "manuscript"
    ts = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    out_name = f"{base}{title_slug}_{ts}.{ext}"
    if _ARGS.output:
        out_dir = Path(_ARGS.output).expanduser().resolve()
    else:
        out_dir = Path(__file__).resolve().parent / "output"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / out_name

    if OUTPUT_FORMAT == "docx":
        export_docx(chapters_out, out_path, title=novel_title, author=author)
    else:
        export_md(chapters_out, out_path, title=novel_title, author=author)

    word_count = sum(len(b.split()) for _, b in chapters_out)
    print(f"\nDone. Manuscript saved to: {out_path}")
    print(f"Approximate word count: {word_count:,}")
    if not TEST_RUN and word_count < 20000:
        print("Note: Target is 20,000–30,000 words. Consider re-running or increasing TARGET_WORDS_PER_CHAPTER.")


if __name__ == "__main__":
    main()
    sys.exit(0)
